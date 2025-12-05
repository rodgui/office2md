"""Create a comprehensive test DOCX with all formatting elements."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document()

# Title
title = doc.add_heading('Documento de Teste Completo', level=1)

# Headings
doc.add_heading('Heading Level 2', level=2)
doc.add_heading('Heading Level 3', level=3)

# Paragraph with formatting
p = doc.add_paragraph('Este parágrafo contém ')
p.add_run('texto em negrito').bold = True
p.add_run(', ')
p.add_run('texto em itálico').italic = True
p.add_run(', e ')
p.add_run('texto sublinhado').underline = True
p.add_run('.')

# Bulleted list
doc.add_paragraph('Item de lista 1', style='List Bullet')
doc.add_paragraph('Item de lista 2', style='List Bullet')
doc.add_paragraph('Item de lista aninhado', style='List Bullet 2')

# Numbered list
doc.add_paragraph('Primeiro item', style='List Number')
doc.add_paragraph('Segundo item', style='List Number')
doc.add_paragraph('Terceiro item', style='List Number')

# Table
doc.add_heading('Tabela de Exemplo', level=2)
table = doc.add_table(rows=3, cols=3)
table.style = 'Light Grid Accent 1'

# Header row
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Coluna 1'
hdr_cells[1].text = 'Coluna 2'
hdr_cells[2].text = 'Coluna 3'

# Data rows
for i in range(1, 3):
    row_cells = table.rows[i].cells
    row_cells[0].text = f'Linha {i} Col 1'
    row_cells[1].text = f'Linha {i} Col 2'
    row_cells[2].text = f'Linha {i} Col 3'

doc.save('examples/input/test_formatting.docx')
print("✅ Created: examples/input/test_formatting.docx")