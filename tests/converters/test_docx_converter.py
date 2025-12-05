"""Tests for DOCX converter."""

import pytest
from pathlib import Path
from unittest.mock import Mock, patch, MagicMock

from office2md.converters.docx_converter import DocxConverter


class TestDocxConverter:
    """Test suite for DocxConverter."""

    @pytest.fixture
    def sample_docx(self, tmp_path):
        """Create a sample DOCX file for testing."""
        from docx import Document
        
        doc_path = tmp_path / "test.docx"
        doc = Document()
        
        # Add heading
        doc.add_heading("Test Document", level=1)
        
        # Add paragraph with formatting
        para = doc.add_paragraph()
        para.add_run("This is ").bold = False
        bold_run = para.add_run("bold text")
        bold_run.bold = True
        para.add_run(" and normal text.")
        
        # Add list
        doc.add_paragraph("Item 1", style='List Bullet')
        doc.add_paragraph("Item 2", style='List Bullet')
        
        # Add table
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Header 1"
        table.cell(0, 1).text = "Header 2"
        table.cell(1, 0).text = "Data 1"
        table.cell(1, 1).text = "Data 2"
        
        doc.save(doc_path)
        return doc_path

    def test_converter_initialization(self, sample_docx):
        """Test converter initializes correctly."""
        converter = DocxConverter(str(sample_docx))
        
        assert converter.input_path == sample_docx
        assert converter.output_path == sample_docx.with_suffix(".md")

    def test_convert_basic_document(self, sample_docx):
        """Test basic document conversion."""
        converter = DocxConverter(str(sample_docx))
        markdown = converter.convert()
        
        # Check heading
        assert "# Test Document" in markdown
        
        # Check bold text
        assert "**bold text**" in markdown
        
        # Check list items
        assert "- Item 1" in markdown
        assert "- Item 2" in markdown

    def test_convert_table(self, sample_docx):
        """Test table conversion."""
        converter = DocxConverter(str(sample_docx))
        markdown = converter.convert()
        
        # Check table structure
        assert "| Header 1 | Header 2 |" in markdown
        assert "| --- | --- |" in markdown
        assert "| Data 1 | Data 2 |" in markdown

    def test_output_path_generation(self, sample_docx):
        """Test output path is generated correctly."""
        converter = DocxConverter(str(sample_docx))
        
        assert converter.output_path.suffix == ".md"
        assert converter.output_path.stem == sample_docx.stem

    def test_custom_output_path(self, sample_docx, tmp_path):
        """Test custom output path."""
        output_path = tmp_path / "custom_output.md"
        converter = DocxConverter(str(sample_docx), str(output_path))
        
        assert converter.output_path == output_path

    def test_save_creates_file(self, sample_docx, tmp_path):
        """Test save() creates markdown file."""
        output_path = tmp_path / "output.md"
        converter = DocxConverter(str(sample_docx), str(output_path))
        converter.save()
        
        assert output_path.exists()
        content = output_path.read_text()
        assert "# Test Document" in content

    def test_images_directory_created(self, sample_docx, tmp_path):
        """Test images directory is created when extract_images=True."""
        output_path = tmp_path / "output.md"
        converter = DocxConverter(
            str(sample_docx), 
            str(output_path),
            extract_images=True
        )
        
        assert converter.images_dir == tmp_path / "output_images"

    def test_skip_images_mode(self, sample_docx):
        """Test skip_images mode."""
        converter = DocxConverter(str(sample_docx), skip_images=True)
        
        assert converter.skip_images is True
        assert converter.extract_images is False

    def test_embed_images_mode(self, sample_docx):
        """Test embed_images mode."""
        converter = DocxConverter(str(sample_docx), embed_images=True)
        
        assert converter.embed_images is True
        assert converter.extract_images is False


class TestDocxConverterBoldFormatting:
    """Test bold formatting fixes."""

    @pytest.fixture
    def converter(self, tmp_path):
        """Create a converter with a dummy file."""
        from docx import Document
        
        doc_path = tmp_path / "bold_test.docx"
        doc = Document()
        doc.add_paragraph("Normal text")
        doc.save(doc_path)
        
        return DocxConverter(str(doc_path))

    def test_fix_multiple_asterisks(self, converter):
        """Test that **** becomes **."""
        input_text = "This is ****bold**** text"
        result = converter._fix_bold_formatting(input_text)
        
        assert "****" not in result
        assert "**bold**" in result

    def test_section_title_detection(self, converter):
        """Test section title detection."""
        # Short, no punctuation = title
        assert converter._is_section_title("Documentos Relacionados") is True
        
        # Long text = not title
        assert converter._is_section_title("A" * 70) is False
        
        # Ends with period = not title
        assert converter._is_section_title("This is a sentence.") is False

    def test_separator_for_title(self, converter):
        """Test separator selection for titles."""
        before_sep, after_sep = converter._get_separator("x", "y", "Title")
        
        # Title should get newlines
        assert "\n" in before_sep
        assert "\n" in after_sep

    def test_separator_for_inline(self, converter):
        """Test separator selection for inline text."""
        before_sep, after_sep = converter._get_separator("x", "y", "This is a long inline text that ends with period.")
        
        # Inline should get spaces, not newlines
        assert before_sep in [" ", ""]
        assert after_sep in [" ", ""]


class TestDocxConverterImages:
    """Test image extraction."""

    @pytest.fixture
    def converter_with_images(self, tmp_path):
        """Create a converter configured for image extraction."""
        from docx import Document
        
        doc_path = tmp_path / "image_test.docx"
        doc = Document()
        doc.add_paragraph("Document with images")
        doc.save(doc_path)
        
        output_path = tmp_path / "output.md"
        return DocxConverter(
            str(doc_path),
            str(output_path),
            extract_images=True
        )

    def test_image_counter_initialization(self, converter_with_images):
        """Test image counter starts at 0."""
        assert converter_with_images._image_index == 0

    def test_get_next_image_ref_empty(self, converter_with_images):
        """Test get_next_image_ref returns empty when no images."""
        ref = converter_with_images._get_next_image_ref()
        assert ref == ""

    def test_get_next_image_ref_with_images(self, converter_with_images):
        """Test get_next_image_ref returns correct reference."""
        converter_with_images._mammoth_images = [
            "![](./output_images/image_1.png)",
            "![](./output_images/image_2.png)",
        ]
        
        ref1 = converter_with_images._get_next_image_ref()
        assert "image_1.png" in ref1
        
        ref2 = converter_with_images._get_next_image_ref()
        assert "image_2.png" in ref2
        
        ref3 = converter_with_images._get_next_image_ref()
        assert ref3 == ""
