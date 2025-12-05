"""Tests for Docling converter."""

import pytest
from pathlib import Path
from unittest.mock import patch, MagicMock


class TestDoclingConverter:
    """Test suite for DoclingConverter."""

    @pytest.fixture
    def sample_docx(self, tmp_path):
        """Create a sample DOCX file."""
        from docx import Document
        
        doc_path = tmp_path / "test.docx"
        doc = Document()
        doc.add_heading("Docling Test", level=1)
        doc.add_paragraph("Test paragraph")
        doc.save(doc_path)
        return doc_path

    @patch('office2md.converters.docling_converter.DOCLING_AVAILABLE', False)
    def test_converter_raises_without_docling(self, sample_docx):
        """Test converter raises error when Docling not available."""
        from office2md.converters.docling_converter import DoclingConverter
        
        with pytest.raises(RuntimeError, match="Docling is not installed"):
            DoclingConverter(str(sample_docx))

    @patch('office2md.converters.docling_converter.DOCLING_AVAILABLE', True)
    @patch('office2md.converters.docling_converter.DocumentConverter')
    def test_convert_calls_docling(self, mock_docling_class, sample_docx):
        """Test convert() calls Docling."""
        # Setup mock
        mock_doc = MagicMock()
        mock_doc.export_to_markdown.return_value = "# Docling Test\n\nTest paragraph"
        
        mock_result = MagicMock()
        mock_result.document = mock_doc
        
        mock_converter = MagicMock()
        mock_converter.convert.return_value = mock_result
        mock_docling_class.return_value = mock_converter
        
        from office2md.converters.docling_converter import DoclingConverter
        
        converter = DoclingConverter(str(sample_docx), skip_images=True)
        result = converter.convert()
        
        # Verify Docling was called
        mock_converter.convert.assert_called_once()
        assert "# Docling Test" in result

    @patch('office2md.converters.docling_converter.DOCLING_AVAILABLE', True)
    @patch('office2md.converters.docling_converter.DocumentConverter')
    def test_convert_handles_error(self, mock_docling_class, sample_docx):
        """Test convert() handles Docling errors."""
        mock_docling_class.return_value.convert.side_effect = Exception("Docling error")
        
        from office2md.converters.docling_converter import DoclingConverter
        
        converter = DoclingConverter(str(sample_docx), skip_images=True)
        
        with pytest.raises(Exception, match="Docling error"):
            converter.convert()