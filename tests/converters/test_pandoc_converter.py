"""Tests for Pandoc converter."""

import pytest
from pathlib import Path
from unittest.mock import patch, MagicMock
import subprocess

from office2md.converters.pandoc_converter import PandocConverter, is_pandoc_available


class TestPandocAvailability:
    """Test Pandoc availability detection."""

    @patch('shutil.which')
    def test_pandoc_available(self, mock_which):
        """Test when Pandoc is installed."""
        mock_which.return_value = "/usr/local/bin/pandoc"
        assert is_pandoc_available() is True

    @patch('shutil.which')
    def test_pandoc_not_available(self, mock_which):
        """Test when Pandoc is not installed."""
        mock_which.return_value = None
        assert is_pandoc_available() is False


class TestPandocConverter:
    """Test suite for PandocConverter."""

    @pytest.fixture
    def sample_docx(self, tmp_path):
        """Create a sample DOCX file."""
        from docx import Document
        
        doc_path = tmp_path / "test.docx"
        doc = Document()
        doc.add_heading("Pandoc Test", level=1)
        doc.add_paragraph("Test paragraph")
        doc.save(doc_path)
        return doc_path

    @patch('office2md.converters.pandoc_converter.PANDOC_AVAILABLE', True)
    def test_converter_initialization(self, sample_docx):
        """Test converter initializes when Pandoc available."""
        converter = PandocConverter(str(sample_docx))
        assert converter.input_path == sample_docx

    @patch('office2md.converters.pandoc_converter.PANDOC_AVAILABLE', False)
    def test_converter_raises_without_pandoc(self, sample_docx):
        """Test converter raises error when Pandoc not available."""
        with pytest.raises(RuntimeError, match="Pandoc is not installed"):
            PandocConverter(str(sample_docx))

    @patch('office2md.converters.pandoc_converter.PANDOC_AVAILABLE', True)
    @patch('subprocess.run')
    def test_convert_calls_pandoc(self, mock_run, sample_docx):
        """Test convert() calls Pandoc subprocess."""
        mock_run.return_value = MagicMock(
            returncode=0,
            stdout="# Pandoc Test\n\nTest paragraph",
            stderr=""
        )
        
        converter = PandocConverter(str(sample_docx), skip_images=True)
        result = converter.convert()
        
        # Verify Pandoc was called
        mock_run.assert_called_once()
        call_args = mock_run.call_args[0][0]
        assert call_args[0] == 'pandoc'
        assert '-f' in call_args
        assert 'docx' in call_args

    @patch('office2md.converters.pandoc_converter.PANDOC_AVAILABLE', True)
    @patch('subprocess.run')
    def test_convert_returns_markdown(self, mock_run, sample_docx):
        """Test convert() returns markdown content."""
        expected_md = "# Heading\n\nParagraph text"
        mock_run.return_value = MagicMock(
            returncode=0,
            stdout=expected_md,
            stderr=""
        )
        
        converter = PandocConverter(str(sample_docx), skip_images=True)
        result = converter.convert()
        
        assert result == expected_md

    @patch('office2md.converters.pandoc_converter.PANDOC_AVAILABLE', True)
    @patch('subprocess.run')
    def test_convert_handles_pandoc_error(self, mock_run, sample_docx):
        """Test convert() handles Pandoc errors."""
        mock_run.return_value = MagicMock(
            returncode=1,
            stdout="",
            stderr="Pandoc error: invalid input"
        )
        
        converter = PandocConverter(str(sample_docx), skip_images=True)
        
        with pytest.raises(RuntimeError, match="Pandoc error"):
            converter.convert()

    @patch('office2md.converters.pandoc_converter.PANDOC_AVAILABLE', True)
    @patch('subprocess.run')
    def test_convert_handles_timeout(self, mock_run, sample_docx):
        """Test convert() handles timeout."""
        mock_run.side_effect = subprocess.TimeoutExpired(cmd='pandoc', timeout=120)
        
        converter = PandocConverter(str(sample_docx), skip_images=True)
        
        with pytest.raises(RuntimeError, match="timed out"):
            converter.convert()