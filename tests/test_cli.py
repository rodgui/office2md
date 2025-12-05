"""Tests for CLI module."""

import pytest
from pathlib import Path
from unittest.mock import patch, MagicMock

from office2md.cli import parse_args, convert_file, batch_convert, main


class TestParseArgs:
    """Test argument parsing."""

    def test_basic_input(self):
        """Test basic input file argument."""
        args = parse_args(["document.docx"])
        assert args.input == "document.docx"

    def test_output_option(self):
        """Test -o/--output option."""
        args = parse_args(["input.docx", "-o", "output.md"])
        assert args.output == "output.md"

    def test_batch_mode(self):
        """Test --batch flag."""
        args = parse_args(["--batch", "./input"])
        assert args.batch is True
        assert args.input == "./input"

    def test_recursive_flag(self):
        """Test --recursive flag."""
        args = parse_args(["--batch", "./input", "-r"])
        assert args.recursive is True

    def test_use_pandoc_flag(self):
        """Test --use-pandoc flag."""
        args = parse_args(["doc.docx", "--use-pandoc"])
        assert args.use_pandoc is True

    def test_use_docling_flag(self):
        """Test --use-docling flag."""
        args = parse_args(["doc.pdf", "--use-docling"])
        assert args.use_docling is True

    def test_no_mammoth_flag(self):
        """Test --no-mammoth flag."""
        args = parse_args(["doc.docx", "--no-mammoth"])
        assert args.no_mammoth is True

    def test_embed_images_flag(self):
        """Test --embed-images flag."""
        args = parse_args(["doc.docx", "--embed-images"])
        assert args.embed_images is True

    def test_skip_images_flag(self):
        """Test --skip-images flag."""
        args = parse_args(["doc.docx", "--skip-images"])
        assert args.skip_images is True

    def test_first_sheet_only_flag(self):
        """Test --first-sheet-only flag."""
        args = parse_args(["data.xlsx", "--first-sheet-only"])
        assert args.first_sheet_only is True

    def test_no_notes_flag(self):
        """Test --no-notes flag."""
        args = parse_args(["slides.pptx", "--no-notes"])
        assert args.no_notes is True

    def test_verbose_flag(self):
        """Test -v/--verbose flag."""
        args = parse_args(["doc.docx", "-v"])
        assert args.verbose is True


class TestConvertFile:
    """Test convert_file function."""

    @pytest.fixture
    def sample_docx(self, tmp_path):
        """Create a sample DOCX file."""
        from docx import Document
        
        doc_path = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("Test content")
        doc.save(doc_path)
        return doc_path

    def test_convert_existing_file(self, sample_docx, tmp_path):
        """Test converting an existing file."""
        output_path = tmp_path / "output.md"
        result = convert_file(str(sample_docx), str(output_path))
        
        assert result is True
        assert output_path.exists()

    def test_convert_nonexistent_file(self, tmp_path):
        """Test converting a nonexistent file."""
        result = convert_file(str(tmp_path / "nonexistent.docx"))
        assert result is False

    @patch('office2md.converters.pandoc_converter.PANDOC_AVAILABLE', True)
    @patch('subprocess.run')
    def test_convert_with_pandoc(self, mock_run, sample_docx, tmp_path):
        """Test converting with --use-pandoc."""
        mock_run.return_value = MagicMock(
            returncode=0,
            stdout="# Test\n\nContent",
            stderr=""
        )
        
        output_path = tmp_path / "output.md"
        result = convert_file(
            str(sample_docx),
            str(output_path),
            use_pandoc=True,
            skip_images=True
        )
        
        assert result is True


class TestBatchConvert:
    """Test batch_convert function."""

    @pytest.fixture
    def batch_input(self, tmp_path):
        """Create a directory with sample files."""
        from docx import Document
        
        input_dir = tmp_path / "input"
        input_dir.mkdir()
        
        # Create sample files
        for i in range(3):
            doc = Document()
            doc.add_paragraph(f"Document {i}")
            doc.save(input_dir / f"doc{i}.docx")
        
        return input_dir

    def test_batch_convert_directory(self, batch_input, tmp_path):
        """Test batch converting a directory."""
        output_dir = tmp_path / "output"
        success, failure = batch_convert(str(batch_input), str(output_dir))
        
        assert success == 3
        assert failure == 0
        assert (output_dir / "doc0.md").exists()

    def test_batch_convert_empty_directory(self, tmp_path):
        """Test batch converting empty directory."""
        empty_dir = tmp_path / "empty"
        empty_dir.mkdir()
        
        success, failure = batch_convert(str(empty_dir))
        assert success == 0
        assert failure == 0

    def test_batch_convert_nonexistent_directory(self, tmp_path):
        """Test batch converting nonexistent directory."""
        success, failure = batch_convert(str(tmp_path / "nonexistent"))
        assert success == 0
        assert failure == 1


class TestMain:
    """Test main CLI function."""

    @pytest.fixture
    def sample_docx(self, tmp_path):
        """Create a sample DOCX file."""
        from docx import Document
        
        doc_path = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("Test")
        doc.save(doc_path)
        return doc_path

    def test_main_no_input(self):
        """Test main with no input returns error."""
        result = main([])
        assert result == 1

    def test_main_single_file(self, sample_docx, tmp_path):
        """Test main with single file."""
        output = tmp_path / "output.md"
        result = main([str(sample_docx), "-o", str(output)])
        
        assert result == 0
        assert output.exists()

    def test_main_batch_mode(self, tmp_path):
        """Test main in batch mode."""
        from docx import Document
        
        input_dir = tmp_path / "input"
        input_dir.mkdir()
        output_dir = tmp_path / "output"
        
        doc = Document()
        doc.add_paragraph("Test")
        doc.save(input_dir / "test.docx")
        
        result = main(["--batch", str(input_dir), "-o", str(output_dir)])
        
        assert result == 0
