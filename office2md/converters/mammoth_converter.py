"""Mammoth-based DOCX converter with enhanced table and image support."""

import logging
import re
from pathlib import Path
from typing import Optional, Dict, Any
from io import BytesIO

from office2md.converters.base_converter import BaseConverter

logger = logging.getLogger(__name__)

try:
    import mammoth
    MAMMOTH_AVAILABLE = True
except ImportError:
    MAMMOTH_AVAILABLE = False

try:
    from markdownify import markdownify as md
    MARKDOWNIFY_AVAILABLE = True
except ImportError:
    MARKDOWNIFY_AVAILABLE = False

try:
    import docx
    from docx.table import Table
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False


class MammothConverter(BaseConverter):
    """
    DOCX converter using Mammoth for HTML extraction and markdownify for conversion.
    
    Uses python-docx for enhanced table extraction when available.
    """

    def __init__(
        self,
        input_path: str,
        output_path: Optional[str] = None,
        **kwargs
    ):
        """Initialize Mammoth converter."""
        super().__init__(input_path, output_path, **kwargs)
        
        if not MAMMOTH_AVAILABLE:
            raise RuntimeError("Mammoth not available. Install with: pip install mammoth")

    def convert(self) -> str:
        """Convert DOCX to Markdown using Mammoth."""
        logger.info(f"Converting with Mammoth: {self.input_path}")
        
        # Create images directory
        if self.images_dir and self.extract_images:
            self.images_dir.mkdir(parents=True, exist_ok=True)
        
        # Convert with Mammoth
        with open(self.input_path, "rb") as docx_file:
            result = mammoth.convert_to_html(
                docx_file,
                convert_image=mammoth.images.img_element(self._handle_image) if self.extract_images else None
            )
        
        html = result.value
        
        # Log any warnings
        for message in result.messages:
            logger.warning(f"Mammoth: {message}")
        
        # Convert HTML to Markdown
        if MARKDOWNIFY_AVAILABLE:
            markdown = md(html, heading_style="ATX", bullets="-")
        else:
            markdown = self._basic_html_to_markdown(html)
        
        # Enhance with python-docx tables if available
        if PYTHON_DOCX_AVAILABLE:
            markdown = self._enhance_tables(markdown)
        
        # Clean up
        markdown = self._cleanup_markdown(markdown)
        
        logger.info("Mammoth conversion completed")
        return markdown

    def _handle_image(self, image) -> Dict[str, Any]:
        """Handle image extraction from Mammoth."""
        if self.skip_images:
            return {}
        
        try:
            with image.open() as image_bytes:
                image_data = image_bytes.read()
            
            ext = image.content_type.split('/')[-1] if image.content_type else 'png'
            if ext == 'jpeg':
                ext = 'jpg'
            
            ref = self._process_image(image_data, ext)
            if ref:
                # Extract path from markdown reference
                match = re.search(r'\(([^)]+)\)', ref)
                if match:
                    return {"src": match.group(1)}
            
        except Exception as e:
            logger.warning(f"Failed to process image: {e}")
        
        return {}

    def _enhance_tables(self, markdown: str) -> str:
        """Enhance tables using python-docx for better extraction."""
        try:
            doc = docx.Document(self.input_path)
            
            if not doc.tables:
                return markdown
            
            # Extract tables from python-docx
            table_markdowns = []
            for table in doc.tables:
                table_md = self._table_to_markdown(table)
                if table_md:
                    table_markdowns.append(table_md)
            
            # If we have better table representations, try to replace
            # This is a heuristic - Mammoth tables may be incomplete
            if table_markdowns:
                # Check if markdown has incomplete tables
                existing_tables = re.findall(r'\|[^\n]+\|(?:\n\|[^\n]+\|)*', markdown)
                
                if len(existing_tables) < len(table_markdowns):
                    # Append missing tables
                    markdown += "\n\n" + "\n\n".join(table_markdowns[len(existing_tables):])
            
            return markdown
            
        except Exception as e:
            logger.debug(f"Table enhancement failed: {e}")
            return markdown

    def _table_to_markdown(self, table) -> str:
        """Convert python-docx table to Markdown."""
        rows = []
        for row in table.rows:
            cells = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
            rows.append(cells)
        
        if not rows:
            return ""
        
        # Normalize column count
        max_cols = max(len(row) for row in rows)
        normalized = [row + [''] * (max_cols - len(row)) for row in rows]
        
        lines = []
        lines.append('| ' + ' | '.join(normalized[0]) + ' |')
        lines.append('| ' + ' | '.join(['---'] * max_cols) + ' |')
        
        for row in normalized[1:]:
            lines.append('| ' + ' | '.join(row) + ' |')
        
        return '\n'.join(lines)

    def _basic_html_to_markdown(self, html: str) -> str:
        """Basic HTML to Markdown conversion when markdownify is not available."""
        # Headers
        html = re.sub(r'<h1[^>]*>(.*?)</h1>', r'# \1\n', html, flags=re.DOTALL)
        html = re.sub(r'<h2[^>]*>(.*?)</h2>', r'## \1\n', html, flags=re.DOTALL)
        html = re.sub(r'<h3[^>]*>(.*?)</h3>', r'### \1\n', html, flags=re.DOTALL)
        html = re.sub(r'<h4[^>]*>(.*?)</h4>', r'#### \1\n', html, flags=re.DOTALL)
        
        # Bold and italic
        html = re.sub(r'<strong[^>]*>(.*?)</strong>', r'**\1**', html, flags=re.DOTALL)
        html = re.sub(r'<b[^>]*>(.*?)</b>', r'**\1**', html, flags=re.DOTALL)
        html = re.sub(r'<em[^>]*>(.*?)</em>', r'*\1*', html, flags=re.DOTALL)
        html = re.sub(r'<i[^>]*>(.*?)</i>', r'*\1*', html, flags=re.DOTALL)
        
        # Lists
        html = re.sub(r'<li[^>]*>(.*?)</li>', r'- \1\n', html, flags=re.DOTALL)
        html = re.sub(r'<[uo]l[^>]*>', '', html)
        html = re.sub(r'</[uo]l>', '\n', html)
        
        # Paragraphs and breaks
        html = re.sub(r'<p[^>]*>(.*?)</p>', r'\1\n\n', html, flags=re.DOTALL)
        html = re.sub(r'<br\s*/?>', '\n', html)
        
        # Images
        html = re.sub(r'<img[^>]*src="([^"]*)"[^>]*/?\s*>', r'![](\1)', html)
        
        # Links
        html = re.sub(r'<a[^>]*href="([^"]*)"[^>]*>(.*?)</a>', r'[\2](\1)', html, flags=re.DOTALL)
        
        # Remove remaining tags
        html = re.sub(r'<[^>]+>', '', html)
        
        # Clean up entities
        html = html.replace('&nbsp;', ' ')
        html = html.replace('&amp;', '&')
        html = html.replace('&lt;', '<')
        html = html.replace('&gt;', '>')
        
        return html

    def _cleanup_markdown(self, markdown: str) -> str:
        """Clean up the generated Markdown."""
        # Remove multiple blank lines
        markdown = re.sub(r'\n{3,}', '\n\n', markdown)
        
        # Clean up trailing whitespace
        lines = markdown.split('\n')
        lines = [line.rstrip() for line in lines]
        markdown = '\n'.join(lines)
        
        return markdown.strip()